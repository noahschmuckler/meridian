#!/usr/bin/env python3
"""Generate a structured Word template from a Meridian module JSON file.

The template uses strict heading styles and field markers so it can be
parsed back to JSON by mammoth.js in the browser. The format is:

  Heading 1: Module title (e.g. "ADHD Stimulants — Inherited Patient")

  Heading 2: Section markers (exact text matters):
    "Introduction", "Checklist Items", "Green Zone",
    "Escalation Items", "Context", "FAQ Reference"

  Heading 3: Item or FAQ entry headers with IDs in brackets:
    Checklist:  "[pdmp] PDMP reviewed — no concerning pattern"
    Escalation: "[dx-vibes] Diagnosis based on self-report alone..."
    FAQ:        "[pdmp] PDMP Review"

  Bold field labels on their own line (parser keys on these):
    "FAQ Title:", "Question:", "Label:", "Narrative:", "SmartPhrase:"

  Instruction paragraphs start with ">>>" and are stripped during parse.
"""

import json
import os
import re
import sys
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Colors matching Meridian ──
GREEN = RGBColor(0x1A, 0x5C, 0x3A)
BROWN = RGBColor(0x7A, 0x3B, 0x1E)
GRAY = RGBColor(0x6B, 0x65, 0x60)
DARK = RGBColor(0x1C, 0x1A, 0x16)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)


def html_to_plain(html):
    """Strip HTML tags, decode entities."""
    text = re.sub(r'<br\s*/?>', '\n', html)
    text = re.sub(r'</p>\s*<p[^>]*>', '\n\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'&amp;', '&', text)
    text = re.sub(r'&lt;', '<', text)
    text = re.sub(r'&gt;', '>', text)
    text = re.sub(r'&quot;', '"', text)
    text = re.sub(r'&#39;', "'", text)
    text = re.sub(r'&nbsp;', ' ', text)
    return text.strip()


def add_instruction(doc, text):
    """Add an instruction paragraph (prefixed with >>> for parser stripping)."""
    p = doc.add_paragraph()
    run = p.add_run('>>> ' + text)
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GRAY
    run.font.italic = True
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)


def add_field(doc, label, value, label_color=DARK):
    """Add a bold-label: value paragraph."""
    p = doc.add_paragraph()
    run_label = p.add_run(label + ': ')
    run_label.font.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = label_color
    run_val = p.add_run(value)
    run_val.font.size = Pt(11)
    run_val.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(6)
    return p


def add_separator(doc):
    """Add a thin horizontal rule via a styled paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('─' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GRAY


def generate_template(module_path, output_path=None):
    """Generate a DOCX template from a module JSON file."""
    with open(module_path, 'r', encoding='utf-8') as f:
        mod = json.load(f)

    if output_path is None:
        output_path = os.path.splitext(module_path)[0] + '.docx'

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Segoe UI'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Segoe UI'
        h_style.font.color.rgb = GREEN if level <= 2 else DARK

    # ════════════════════════════════════════
    # MODULE TITLE (Heading 1)
    # ════════════════════════════════════════
    doc.add_heading(mod['default_title'], level=1)

    add_instruction(doc,
        'TEMPLATE FORMAT RULES — Do not delete or reorder section headings. '
        'Lines starting with ">>>" are instructions and will be stripped on upload. '
        'Item IDs in [brackets] must be unique within the module. '
        'Keep IDs lowercase, hyphenated, no spaces (e.g., [pdmp], [high-dose]).')

    # ════════════════════════════════════════
    # INTRODUCTION (Heading 2)
    # ════════════════════════════════════════
    doc.add_heading('Introduction', level=2)
    add_instruction(doc,
        'One paragraph of introductory context for the clinician. '
        'This appears at the top of the module in italics.')
    doc.add_paragraph(mod.get('landing_intro', ''))

    # ════════════════════════════════════════
    # CHECKLIST ITEMS (Heading 2)
    # ════════════════════════════════════════
    doc.add_heading('Checklist Items', level=2)
    add_instruction(doc,
        'Exactly 4 items required. Each item is a Heading 3 with the format: '
        '[item-id] Statement text. The item-id must match a FAQ entry ID below.')

    for item in mod.get('checklist', []):
        doc.add_heading(
            '[' + item['item_id'] + '] ' + item['statement'],
            level=3
        )

    if len(mod.get('checklist', [])) < 4:
        for i in range(4 - len(mod.get('checklist', []))):
            add_instruction(doc, 'Add checklist item ' + str(len(mod.get('checklist', [])) + i + 1) + ' here as Heading 3: [item-id] Statement text')

    # ════════════════════════════════════════
    # GREEN ZONE (Heading 2)
    # ════════════════════════════════════════
    doc.add_heading('Green Zone', level=2)
    add_instruction(doc,
        'The success state shown when all checklist items are checked. '
        'Three fields required: Label, Narrative, SmartPhrase.')

    gz = mod.get('green_zone', {}) or {}
    add_field(doc, 'Label', gz.get('zone_label', ''))
    add_field(doc, 'Narrative', html_to_plain(gz.get('narrative_html', '')))
    add_field(doc, 'SmartPhrase', gz.get('smartphrase', ''))

    # ════════════════════════════════════════
    # ESCALATION ITEMS (Heading 2)
    # ════════════════════════════════════════
    doc.add_heading('Escalation Items', level=2)
    add_instruction(doc,
        'Up to 10 items. Each item is a Heading 3 with the format: '
        '[item-id] Statement text. The item-id must match a FAQ entry ID below.')

    for item in mod.get('escalation', []):
        doc.add_heading(
            '[' + item['item_id'] + '] ' + item['statement'],
            level=3
        )

    # ════════════════════════════════════════
    # CONTEXT (Heading 2) — optional
    # ════════════════════════════════════════
    doc.add_heading('Context', level=2)
    add_instruction(doc,
        'Optional. A brief contextual note displayed at the bottom of the module. '
        'Delete this section if not needed. Two fields: Label and Text.')

    cs = mod.get('context_strip')
    if cs:
        add_field(doc, 'Label', cs.get('label', ''))
        add_field(doc, 'Text', cs.get('text', ''))
    else:
        add_field(doc, 'Label', '')
        add_field(doc, 'Text', '')

    # ════════════════════════════════════════
    # FOOTER (Heading 2)
    # ════════════════════════════════════════
    doc.add_heading('Footer', level=2)
    add_instruction(doc, 'One line of footer text displayed at the bottom of the module.')
    doc.add_paragraph(mod.get('footer_note', ''))

    # ════════════════════════════════════════
    # FAQ REFERENCE (Heading 2)
    # ════════════════════════════════════════
    doc.add_heading('FAQ Reference', level=2)
    add_instruction(doc,
        'One entry per checklist or escalation item. Each entry starts with a '
        'Heading 3: [item-id] Topic Name. The item-id must match an item above. '
        'Then "FAQ Title:" followed by the display title. '
        'Then pairs of "Question:" and answer paragraphs. '
        'Answer text can use bold and italic formatting — it will be preserved.')

    for faq in mod.get('faqs', []):
        add_separator(doc)
        doc.add_heading(
            '[' + faq['faq_id'] + '] ' + faq['topic'],
            level=3
        )
        add_field(doc, 'FAQ Title', faq.get('title', ''))

        for qa in faq.get('items', []):
            doc.add_paragraph()  # spacer
            add_field(doc, 'Question', qa.get('question', ''))

            # Answer — preserve as plain text (formatting via Word)
            answer_text = html_to_plain(qa.get('answer_html', ''))
            p = doc.add_paragraph()
            run = p.add_run(answer_text)
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(8)

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 2:
        # Default: generate all modules
        modules_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'modules')
        index_path = os.path.join(modules_dir, 'index.json')
        with open(index_path) as f:
            index = json.load(f)

        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
        os.makedirs(output_dir, exist_ok=True)

        for entry in index['modules']:
            mod_path = os.path.join(modules_dir, entry['file'])
            out_path = os.path.join(output_dir, entry['module_id'] + '.docx')
            print(f"Generating {entry['module_id']}...")
            generate_template(mod_path, out_path)
            print(f"  -> {out_path}")

        print(f"\nDone. {len(index['modules'])} templates generated.")
    else:
        module_path = sys.argv[1]
        output_path = sys.argv[2] if len(sys.argv) > 2 else None
        result = generate_template(module_path, output_path)
        print(f"Generated: {result}")


if __name__ == '__main__':
    main()
